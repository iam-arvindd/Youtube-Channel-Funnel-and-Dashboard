from dotenv import load_dotenv
from pathlib import Path

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

import os
import uuid
import logging
from datetime import datetime, timezone, timedelta
from typing import List, Optional, Literal

import bcrypt
import jwt
from cryptography.fernet import Fernet
import io
import asyncio
import base64
import resend
import requests
from googleapiclient.discovery import build as gbuild
from google_auth_oauthlib.flow import Flow
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request as GoogleAuthRequest
from fastapi import FastAPI, APIRouter, Depends, HTTPException, Request, status, UploadFile, File, Query, Header, Response
from fastapi.responses import StreamingResponse
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
from pydantic import BaseModel, Field, ConfigDict
from anthropic import Anthropic

# ---------- Config ----------
MONGO_URL = os.environ['MONGO_URL']
DB_NAME = os.environ['DB_NAME']
JWT_SECRET = os.environ['JWT_SECRET']
ADMIN_EMAIL = os.environ['ADMIN_EMAIL'].lower()
ADMIN_PASSWORD = os.environ['ADMIN_PASSWORD']
ENCRYPTION_KEY = os.environ['API_KEY_ENCRYPTION_KEY']
EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY', '')
STORAGE_URL = "https://integrations.emergentagent.com/objstore/api/v1/storage"
APP_NAME = "wealth-studio"
storage_key_cache = None
JWT_ALGORITHM = "HS256"
fernet = Fernet(ENCRYPTION_KEY.encode() if isinstance(ENCRYPTION_KEY, str) else ENCRYPTION_KEY)

# ---------- Logging ----------
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ---------- DB ----------
mongo_client = AsyncIOMotorClient(MONGO_URL)
db = mongo_client[DB_NAME]

# ---------- App ----------
app = FastAPI(title="Finance YT Command Center")
api = APIRouter(prefix="/api")

# ---------- Helpers ----------
def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()

def hash_password(pw: str) -> str:
    return bcrypt.hashpw(pw.encode(), bcrypt.gensalt()).decode()

def verify_password(pw: str, hashed: str) -> bool:
    try:
        return bcrypt.checkpw(pw.encode(), hashed.encode())
    except Exception:
        return False

def create_token(user_id: str, email: str) -> str:
    payload = {
        "sub": user_id,
        "email": email,
        "exp": datetime.now(timezone.utc) + timedelta(days=7),
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

def encrypt_key(plain: str) -> str:
    return fernet.encrypt(plain.encode()).decode()

def decrypt_key(cipher: str) -> str:
    return fernet.decrypt(cipher.encode()).decode()

def init_storage():
    global storage_key_cache
    if storage_key_cache:
        return storage_key_cache
    if not EMERGENT_LLM_KEY:
        raise RuntimeError("EMERGENT_LLM_KEY not configured")
    r = requests.post(f"{STORAGE_URL}/init", json={"emergent_key": EMERGENT_LLM_KEY}, timeout=30)
    r.raise_for_status()
    storage_key_cache = r.json()["storage_key"]
    return storage_key_cache

def put_object(path: str, data: bytes, content_type: str) -> dict:
    key = init_storage()
    r = requests.put(
        f"{STORAGE_URL}/objects/{path}",
        headers={"X-Storage-Key": key, "Content-Type": content_type},
        data=data, timeout=120,
    )
    r.raise_for_status()
    return r.json()

def get_object(path: str):
    key = init_storage()
    r = requests.get(f"{STORAGE_URL}/objects/{path}", headers={"X-Storage-Key": key}, timeout=60)
    r.raise_for_status()
    return r.content, r.headers.get("Content-Type", "application/octet-stream")

async def get_current_user(request: Request) -> dict:
    auth = request.headers.get("Authorization", "")
    if not auth.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Not authenticated")
    token = auth[7:]
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user = await db.users.find_one({"id": payload["sub"]})
        if not user:
            raise HTTPException(status_code=401, detail="User not found")
        user.pop("_id", None)
        user.pop("password_hash", None)
        return user
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

# ---------- Models ----------
class LoginRequest(BaseModel):
    email: str
    password: str

class ApiKeyRequest(BaseModel):
    api_key: str

class Idea(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    sub_niche: str  # personal_finance | investing | psychology | case_study | side_hustle
    hook: Optional[str] = ""
    rating: int = 5  # 1-10
    tags: List[str] = []
    notes: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

class IdeaCreate(BaseModel):
    title: str
    sub_niche: str
    hook: Optional[str] = ""
    rating: int = 5
    tags: List[str] = []
    notes: Optional[str] = ""

class IdeaUpdate(BaseModel):
    title: Optional[str] = None
    sub_niche: Optional[str] = None
    hook: Optional[str] = None
    rating: Optional[int] = None
    tags: Optional[List[str]] = None
    notes: Optional[str] = None

VideoStage = Literal["idea", "script", "voiceover", "video", "thumbnail", "scheduled", "published"]

class Video(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    sub_niche: str = "investing"
    stage: VideoStage = "idea"
    hook: Optional[str] = ""
    script: Optional[str] = ""
    voiceover_url: Optional[str] = ""
    video_url: Optional[str] = ""
    thumbnail_url: Optional[str] = ""
    youtube_url: Optional[str] = ""
    scheduled_date: Optional[str] = None  # ISO date
    published_date: Optional[str] = None
    notes: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)
    updated_at: str = Field(default_factory=now_iso)

class VideoCreate(BaseModel):
    title: str
    sub_niche: Optional[str] = "investing"
    stage: Optional[VideoStage] = "idea"
    hook: Optional[str] = ""

class VideoUpdate(BaseModel):
    title: Optional[str] = None
    sub_niche: Optional[str] = None
    stage: Optional[VideoStage] = None
    hook: Optional[str] = None
    script: Optional[str] = None
    voiceover_url: Optional[str] = None
    video_url: Optional[str] = None
    thumbnail_url: Optional[str] = None
    youtube_url: Optional[str] = None
    scheduled_date: Optional[str] = None
    published_date: Optional[str] = None
    notes: Optional[str] = None

class AnalyticsEntry(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    video_id: Optional[str] = None
    video_title: str
    date: str  # ISO date
    views: int = 0
    ctr: float = 0.0  # percent
    retention: float = 0.0  # percent
    watch_hours: float = 0.0
    adsense_earnings: float = 0.0  # in INR
    subscribers_gained: int = 0
    rpm: float = 0.0
    created_at: str = Field(default_factory=now_iso)

class AnalyticsCreate(BaseModel):
    video_id: Optional[str] = None
    video_title: str
    date: str
    views: int = 0
    ctr: float = 0.0
    retention: float = 0.0
    watch_hours: float = 0.0
    adsense_earnings: float = 0.0
    subscribers_gained: int = 0
    rpm: float = 0.0

class Affiliate(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    partner: str  # Groww / Zerodha / Upstox / Amazon / Other
    month: str  # YYYY-MM
    clicks: int = 0
    conversions: int = 0
    earnings: float = 0.0  # INR
    notes: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

class AffiliateCreate(BaseModel):
    partner: str
    month: str
    clicks: int = 0
    conversions: int = 0
    earnings: float = 0.0
    notes: Optional[str] = ""

class ChatMessage(BaseModel):
    role: Literal["user", "assistant"]
    content: str

class ChatRequest(BaseModel):
    video_id: Optional[str] = None
    messages: List[ChatMessage]
    system: Optional[str] = None

# ---------- Auth ----------
@api.post("/auth/login")
async def login(body: LoginRequest):
    email = body.email.strip().lower()
    user = await db.users.find_one({"email": email})
    if not user or not verify_password(body.password, user["password_hash"]):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    token = create_token(user["id"], user["email"])
    return {"token": token, "user": {"id": user["id"], "email": user["email"], "name": user.get("name", "Admin")}}

@api.get("/auth/me")
async def me(user: dict = Depends(get_current_user)):
    return user

# ---------- Settings (Anthropic Key) ----------
@api.post("/settings/anthropic-key")
async def save_key(body: ApiKeyRequest, user: dict = Depends(get_current_user)):
    key = body.api_key.strip()
    if len(key) < 20 or not key.startswith("sk-"):
        raise HTTPException(status_code=400, detail="API key looks invalid (must start with 'sk-')")
    encrypted = encrypt_key(key)
    await db.user_settings.update_one(
        {"user_id": user["id"]},
        {"$set": {"anthropic_key_encrypted": encrypted, "updated_at": now_iso()}},
        upsert=True,
    )
    return {"status": "ok", "preview": f"{key[:7]}...{key[-4:]}"}

@api.get("/settings/anthropic-key")
async def get_key_info(user: dict = Depends(get_current_user)):
    doc = await db.user_settings.find_one({"user_id": user["id"]})
    if not doc or "anthropic_key_encrypted" not in doc:
        return {"configured": False}
    try:
        plain = decrypt_key(doc["anthropic_key_encrypted"])
        return {"configured": True, "preview": f"{plain[:7]}...{plain[-4:]}", "updated_at": doc.get("updated_at")}
    except Exception:
        return {"configured": False}

@api.delete("/settings/anthropic-key")
async def delete_key(user: dict = Depends(get_current_user)):
    await db.user_settings.update_one({"user_id": user["id"]}, {"$unset": {"anthropic_key_encrypted": ""}})
    return {"status": "removed"}

# ---------- Generic 3rd-party key storage (youtube, resend) ----------
ALLOWED_KEY_TYPES = {"youtube", "resend", "yt_oauth_client_id", "yt_oauth_client_secret"}

class GenericKeyRequest(BaseModel):
    api_key: str

@api.post("/settings/keys/{key_type}")
async def save_generic_key(key_type: str, body: GenericKeyRequest, user: dict = Depends(get_current_user)):
    if key_type not in ALLOWED_KEY_TYPES:
        raise HTTPException(400, "Unknown key type")
    val = body.api_key.strip()
    if len(val) < 10:
        raise HTTPException(400, "Key looks too short")
    encrypted = encrypt_key(val)
    await db.user_settings.update_one(
        {"user_id": user["id"]},
        {"$set": {f"{key_type}_key_encrypted": encrypted, "updated_at": now_iso()}},
        upsert=True,
    )
    return {"status": "ok", "preview": f"{val[:6]}...{val[-4:]}"}

@api.get("/settings/keys/{key_type}")
async def get_generic_key_info(key_type: str, user: dict = Depends(get_current_user)):
    if key_type not in ALLOWED_KEY_TYPES:
        raise HTTPException(400, "Unknown key type")
    doc = await db.user_settings.find_one({"user_id": user["id"]})
    field = f"{key_type}_key_encrypted"
    if not doc or field not in doc:
        return {"configured": False}
    try:
        plain = decrypt_key(doc[field])
        return {"configured": True, "preview": f"{plain[:6]}...{plain[-4:]}"}
    except Exception:
        return {"configured": False}

@api.delete("/settings/keys/{key_type}")
async def delete_generic_key(key_type: str, user: dict = Depends(get_current_user)):
    if key_type not in ALLOWED_KEY_TYPES:
        raise HTTPException(400, "Unknown key type")
    await db.user_settings.update_one({"user_id": user["id"]}, {"$unset": {f"{key_type}_key_encrypted": ""}})
    return {"status": "removed"}

async def get_user_key(user_id: str, key_type: str) -> Optional[str]:
    doc = await db.user_settings.find_one({"user_id": user_id})
    field = f"{key_type}_key_encrypted"
    if not doc or field not in doc:
        return None
    try:
        return decrypt_key(doc[field])
    except Exception:
        return None

# ---------- Digest config ----------
class DigestConfig(BaseModel):
    enabled: bool = False
    email: Optional[str] = ""
    last_sent_at: Optional[str] = None

@api.get("/settings/digest")
async def get_digest(user: dict = Depends(get_current_user)):
    doc = await db.user_settings.find_one({"user_id": user["id"]})
    d = (doc or {}).get("digest", {})
    return {"enabled": d.get("enabled", False), "email": d.get("email", ""), "last_sent_at": d.get("last_sent_at")}

@api.post("/settings/digest")
async def set_digest(body: DigestConfig, user: dict = Depends(get_current_user)):
    if body.enabled and not body.email:
        raise HTTPException(400, "Email required when digest is enabled")
    await db.user_settings.update_one(
        {"user_id": user["id"]},
        {"$set": {"digest": {"enabled": body.enabled, "email": body.email, "last_sent_at": body.last_sent_at}}},
        upsert=True,
    )
    return {"status": "ok"}

def _build_digest_html(summary: dict, top_videos: list, affiliates_by_partner: dict) -> str:
    rows = ""
    for v in top_videos[:5]:
        rows += f"<tr><td style='padding:8px;border-bottom:1px solid #eee;'>{v.get('video_title','')}</td><td style='padding:8px;border-bottom:1px solid #eee;text-align:right;'>{v.get('views',0):,}</td><td style='padding:8px;border-bottom:1px solid #eee;text-align:right;'>₹{v.get('adsense_earnings',0)}</td></tr>"
    aff_rows = ""
    for partner, amt in affiliates_by_partner.items():
        aff_rows += f"<tr><td style='padding:6px'>{partner}</td><td style='padding:6px;text-align:right'>₹{amt:,.0f}</td></tr>"
    return f"""
    <table width='100%' style='font-family:Helvetica,Arial,sans-serif;color:#0A0A0A;max-width:600px;margin:0 auto;'>
      <tr><td style='padding:24px 0;'>
        <h1 style='color:#00594C;margin:0;font-size:28px;'>Your Weekly Wealth Studio Digest</h1>
        <p style='color:#5C5C5C'>Numbers that matter, in one glance.</p>
      </td></tr>
      <tr><td>
        <table width='100%' cellspacing='0'>
          <tr>
            <td style='background:#E5F2F0;padding:18px;border-radius:10px;'>
              <div style='font-size:11px;color:#5C5C5C;text-transform:uppercase;letter-spacing:.1em'>Total Earnings</div>
              <div style='font-size:28px;font-weight:bold;color:#00594C'>₹{summary.get('total_earnings',0):,.0f}</div>
            </td>
            <td width='12'></td>
            <td style='background:#FCEEEA;padding:18px;border-radius:10px;'>
              <div style='font-size:11px;color:#5C5C5C;text-transform:uppercase;letter-spacing:.1em'>Total Views</div>
              <div style='font-size:28px;font-weight:bold;color:#FF6B4A'>{summary.get('total_views',0):,}</div>
            </td>
          </tr>
        </table>
      </td></tr>
      <tr><td style='padding-top:24px'>
        <h3 style='color:#0A0A0A'>Top videos</h3>
        <table width='100%' style='border-collapse:collapse;font-size:13px;'>{rows or "<tr><td>No analytics logged yet.</td></tr>"}</table>
      </td></tr>
      <tr><td style='padding-top:24px'>
        <h3 style='color:#0A0A0A'>Affiliate income</h3>
        <table width='100%' style='border-collapse:collapse;font-size:13px;'>{aff_rows or "<tr><td>No affiliate entries.</td></tr>"}</table>
      </td></tr>
      <tr><td style='padding:30px 0;color:#8A8A8A;font-size:11px;text-align:center'>
        Wealth Studio · Command Center · This is your weekly digest.
      </td></tr>
    </table>
    """

@api.post("/settings/digest/send-now")
async def send_digest_now(user: dict = Depends(get_current_user)):
    doc = await db.user_settings.find_one({"user_id": user["id"]})
    d = (doc or {}).get("digest", {})
    if not d.get("email"):
        raise HTTPException(400, "Set a digest email first")
    resend_key = await get_user_key(user["id"], "resend")
    if not resend_key:
        raise HTTPException(400, "Add your Resend API key in Settings first")

    # Build summary inline (reuse logic)
    videos = await db.videos.find({}, {"_id": 0}).to_list(1000)
    analytics = await db.analytics.find({}, {"_id": 0}).to_list(2000)
    affs = await db.affiliates.find({}, {"_id": 0}).to_list(2000)
    total_views = sum(a.get("views", 0) for a in analytics)
    total_adsense = sum(a.get("adsense_earnings", 0) for a in analytics)
    total_affiliate = sum(a.get("earnings", 0) for a in affs)
    by_partner = {}
    for a in affs:
        by_partner[a["partner"]] = by_partner.get(a["partner"], 0) + a.get("earnings", 0)
    summary = {"total_earnings": total_adsense + total_affiliate, "total_views": total_views}
    top_videos = sorted(analytics, key=lambda x: x.get("views", 0), reverse=True)

    html = _build_digest_html(summary, top_videos, by_partner)
    resend.api_key = resend_key
    try:
        email = await asyncio.to_thread(resend.Emails.send, {
            "from": "Wealth Studio <onboarding@resend.dev>",
            "to": [d["email"]],
            "subject": "📈 Your Weekly Wealth Studio Digest",
            "html": html,
        })
    except Exception as e:
        raise HTTPException(500, f"Failed to send: {str(e)}")
    await db.user_settings.update_one(
        {"user_id": user["id"]},
        {"$set": {"digest.last_sent_at": now_iso()}},
    )
    return {"status": "sent", "id": email.get("id") if isinstance(email, dict) else None}

# ---------- YouTube auto-fetch ----------
def _extract_video_id(url: str) -> Optional[str]:
    if not url:
        return None
    import re
    patterns = [r"youtu\.be/([A-Za-z0-9_-]{11})", r"v=([A-Za-z0-9_-]{11})", r"shorts/([A-Za-z0-9_-]{11})"]
    for p in patterns:
        m = re.search(p, url)
        if m:
            return m.group(1)
    return None

@api.post("/videos/{vid}/youtube/sync")
async def sync_youtube(vid: str, user: dict = Depends(get_current_user)):
    video = await db.videos.find_one({"id": vid}, {"_id": 0})
    if not video:
        raise HTTPException(404, "Video not found")
    yt_url = video.get("youtube_url")
    yt_id = _extract_video_id(yt_url or "")
    if not yt_id:
        raise HTTPException(400, "Set a valid YouTube URL on this video card first")
    yt_key = await get_user_key(user["id"], "youtube")
    if not yt_key:
        raise HTTPException(400, "Add your YouTube Data API key in Settings first")
    result = await _yt_sync_one(yt_key, vid, video, yt_id)
    return {"status": "ok", **result}

async def _yt_sync_one(yt_key: str, vid: str, video: dict, yt_id: str) -> dict:
    yt = gbuild("youtube", "v3", developerKey=yt_key, cache_discovery=False)
    resp = await asyncio.to_thread(
        lambda: yt.videos().list(part="statistics,snippet,contentDetails", id=yt_id).execute()
    )
    items = resp.get("items", [])
    if not items:
        return {"video_id": vid, "skipped": "not-on-youtube"}
    stats = items[0]["statistics"]
    snippet = items[0]["snippet"]
    views = int(stats.get("viewCount", 0))
    likes = int(stats.get("likeCount", 0))
    comments = int(stats.get("commentCount", 0))
    today = datetime.now(timezone.utc).date().isoformat()
    entry = AnalyticsEntry(
        video_id=vid, video_title=snippet.get("title", video["title"]),
        date=today, views=views, subscribers_gained=0,
    )
    await db.analytics.delete_many({"video_id": vid, "date": today})
    await db.analytics.insert_one(entry.model_dump())
    return {"video_id": vid, "views": views, "likes": likes, "comments": comments, "title": snippet.get("title")}

@api.post("/youtube/sync-all")
async def sync_all(user: dict = Depends(get_current_user)):
    yt_key = await get_user_key(user["id"], "youtube")
    if not yt_key:
        raise HTTPException(400, "Add your YouTube Data API key in Settings first")
    videos = await db.videos.find({"youtube_url": {"$nin": [None, ""]}}, {"_id": 0}).to_list(500)
    results = []
    for v in videos:
        yt_id = _extract_video_id(v.get("youtube_url") or "")
        if not yt_id:
            results.append({"video_id": v["id"], "skipped": "bad-url"}); continue
        try:
            r = await _yt_sync_one(yt_key, v["id"], v, yt_id)
            results.append(r)
        except Exception as e:
            results.append({"video_id": v["id"], "error": str(e)[:120]})
    return {"synced": len([r for r in results if "views" in r]), "total": len(videos), "results": results}

# ---------- Gemini Nano Banana thumbnail generation ----------
class ThumbnailGenRequest(BaseModel):
    prompt: str

@api.post("/videos/{vid}/thumbnail/generate")
async def generate_thumbnail(vid: str, body: ThumbnailGenRequest, user: dict = Depends(get_current_user)):
    if not EMERGENT_LLM_KEY:
        raise HTTPException(500, "EMERGENT_LLM_KEY not configured on server")
    video = await db.videos.find_one({"id": vid})
    if not video:
        raise HTTPException(404, "Video not found")
    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
    except Exception as e:
        raise HTTPException(500, f"emergentintegrations not installed: {e}")
    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=f"thumb-{vid}-{uuid.uuid4()}",
        system_message="You are an expert YouTube thumbnail designer for a finance education channel.",
    )
    chat.with_model("gemini", "gemini-3.1-flash-image-preview").with_params(modalities=["image", "text"])
    full_prompt = (
        f"Create a bold, eye-catching YouTube thumbnail (16:9, 1280x720) for a finance education video titled "
        f"'{video['title']}'. {body.prompt}. Use high contrast colors, large bold text (3-5 words max), "
        f"a striking finance visual (money, charts, shocked expression, or India-relevant imagery). "
        f"Style: modern, professional, finance YouTube channel. No watermarks."
    )
    try:
        text, images = await chat.send_message_multimodal_response(UserMessage(text=full_prompt))
    except Exception as e:
        raise HTTPException(500, f"Image gen failed: {e}")
    if not images:
        raise HTTPException(500, "No image returned by Gemini")
    img = images[0]
    image_bytes = base64.b64decode(img["data"])
    ext = "png"
    path = f"{APP_NAME}/uploads/{user['id']}/{vid}/thumbnail-gen-{uuid.uuid4()}.{ext}"
    try:
        result = put_object(path, image_bytes, "image/png")
    except Exception as e:
        raise HTTPException(500, f"Storage failed: {e}")
    public_path = f"/api/files/{result['path']}"
    await db.files.insert_one({
        "id": str(uuid.uuid4()), "video_id": vid, "kind": "thumbnail",
        "variant": "ai", "ctr": 0.0, "impressions": 0, "clicks": 0,
        "storage_path": result["path"], "original_filename": "ai-thumbnail.png",
        "content_type": "image/png", "size": len(image_bytes),
        "is_deleted": False, "created_at": now_iso(),
    })
    await db.videos.update_one({"id": vid}, {"$set": {"thumbnail_url": public_path, "updated_at": now_iso()}})
    return {"url": public_path}

# ---------- Ideas ----------
@api.get("/ideas")
async def list_ideas(user: dict = Depends(get_current_user)):
    docs = await db.ideas.find({}, {"_id": 0}).sort("created_at", -1).to_list(500)
    return docs

@api.post("/ideas")
async def create_idea(body: IdeaCreate, user: dict = Depends(get_current_user)):
    idea = Idea(**body.model_dump())
    await db.ideas.insert_one(idea.model_dump())
    return idea

@api.patch("/ideas/{idea_id}")
async def update_idea(idea_id: str, body: IdeaUpdate, user: dict = Depends(get_current_user)):
    updates = {k: v for k, v in body.model_dump().items() if v is not None}
    if updates:
        await db.ideas.update_one({"id": idea_id}, {"$set": updates})
    doc = await db.ideas.find_one({"id": idea_id}, {"_id": 0})
    if not doc:
        raise HTTPException(404, "Idea not found")
    return doc

@api.delete("/ideas/{idea_id}")
async def delete_idea(idea_id: str, user: dict = Depends(get_current_user)):
    await db.ideas.delete_one({"id": idea_id})
    return {"status": "deleted"}

@api.post("/ideas/{idea_id}/promote")
async def promote_idea(idea_id: str, user: dict = Depends(get_current_user)):
    """Promote an idea into the pipeline as a Video card (stage=script)."""
    idea = await db.ideas.find_one({"id": idea_id}, {"_id": 0})
    if not idea:
        raise HTTPException(404, "Idea not found")
    video = Video(
        title=idea["title"],
        sub_niche=idea.get("sub_niche", "investing"),
        hook=idea.get("hook", ""),
        stage="script",
    )
    await db.videos.insert_one(video.model_dump())
    return video

# ---------- Videos / Pipeline ----------
@api.get("/videos")
async def list_videos(user: dict = Depends(get_current_user)):
    return await db.videos.find({}, {"_id": 0}).sort("created_at", -1).to_list(500)

@api.post("/videos")
async def create_video(body: VideoCreate, user: dict = Depends(get_current_user)):
    video = Video(**body.model_dump())
    await db.videos.insert_one(video.model_dump())
    return video

@api.get("/videos/{vid}")
async def get_video(vid: str, user: dict = Depends(get_current_user)):
    doc = await db.videos.find_one({"id": vid}, {"_id": 0})
    if not doc:
        raise HTTPException(404, "Video not found")
    return doc

@api.patch("/videos/{vid}")
async def update_video(vid: str, body: VideoUpdate, user: dict = Depends(get_current_user)):
    updates = {k: v for k, v in body.model_dump().items() if v is not None}
    updates["updated_at"] = now_iso()
    await db.videos.update_one({"id": vid}, {"$set": updates})
    doc = await db.videos.find_one({"id": vid}, {"_id": 0})
    if not doc:
        raise HTTPException(404, "Video not found")
    return doc

@api.delete("/videos/{vid}")
async def delete_video(vid: str, user: dict = Depends(get_current_user)):
    await db.videos.delete_one({"id": vid})
    await db.chat_messages.delete_many({"video_id": vid})
    return {"status": "deleted"}

# ---------- Chat history per video ----------
@api.get("/videos/{vid}/messages")
async def get_messages(vid: str, user: dict = Depends(get_current_user)):
    msgs = await db.chat_messages.find({"video_id": vid}, {"_id": 0}).sort("created_at", 1).to_list(500)
    return msgs

@api.delete("/videos/{vid}/messages")
async def clear_messages(vid: str, user: dict = Depends(get_current_user)):
    await db.chat_messages.delete_many({"video_id": vid})
    return {"status": "cleared"}

# ---------- Analytics ----------
@api.get("/analytics")
async def list_analytics(user: dict = Depends(get_current_user)):
    return await db.analytics.find({}, {"_id": 0}).sort("date", -1).to_list(1000)

@api.post("/analytics")
async def create_analytics(body: AnalyticsCreate, user: dict = Depends(get_current_user)):
    entry = AnalyticsEntry(**body.model_dump())
    await db.analytics.insert_one(entry.model_dump())
    return entry

@api.delete("/analytics/{eid}")
async def delete_analytics(eid: str, user: dict = Depends(get_current_user)):
    await db.analytics.delete_one({"id": eid})
    return {"status": "deleted"}

# ---------- Affiliates ----------
@api.get("/affiliates")
async def list_affiliates(user: dict = Depends(get_current_user)):
    return await db.affiliates.find({}, {"_id": 0}).sort("month", -1).to_list(1000)

@api.post("/affiliates")
async def create_affiliate(body: AffiliateCreate, user: dict = Depends(get_current_user)):
    aff = Affiliate(**body.model_dump())
    await db.affiliates.insert_one(aff.model_dump())
    return aff

@api.delete("/affiliates/{aid}")
async def delete_affiliate(aid: str, user: dict = Depends(get_current_user)):
    await db.affiliates.delete_one({"id": aid})
    return {"status": "deleted"}

# ---------- Dashboard summary ----------
@api.get("/dashboard/summary")
async def dashboard_summary(user: dict = Depends(get_current_user)):
    videos = await db.videos.find({}, {"_id": 0}).to_list(1000)
    analytics = await db.analytics.find({}, {"_id": 0}).to_list(2000)
    affiliates = await db.affiliates.find({}, {"_id": 0}).to_list(2000)

    by_stage = {}
    for v in videos:
        by_stage[v["stage"]] = by_stage.get(v["stage"], 0) + 1

    total_views = sum(a.get("views", 0) for a in analytics)
    total_adsense = sum(a.get("adsense_earnings", 0) for a in analytics)
    total_subs = sum(a.get("subscribers_gained", 0) for a in analytics)
    total_affiliate = sum(a.get("earnings", 0) for a in affiliates)
    avg_ctr = (sum(a.get("ctr", 0) for a in analytics) / len(analytics)) if analytics else 0
    avg_retention = (sum(a.get("retention", 0) for a in analytics) / len(analytics)) if analytics else 0

    # last 30 days views trend (group by date)
    trend_map = {}
    for a in analytics:
        d = (a.get("date") or "")[:10]
        if not d:
            continue
        trend_map[d] = trend_map.get(d, 0) + a.get("views", 0)
    trend = [{"date": k, "views": v} for k, v in sorted(trend_map.items())][-30:]

    return {
        "total_videos": len(videos),
        "by_stage": by_stage,
        "total_views": total_views,
        "total_adsense": round(total_adsense, 2),
        "total_affiliate": round(total_affiliate, 2),
        "total_earnings": round(total_adsense + total_affiliate, 2),
        "total_subscribers_gained": total_subs,
        "avg_ctr": round(avg_ctr, 2),
        "avg_retention": round(avg_retention, 2),
        "trend": trend,
    }

# ---------- Claude Chat (streaming) ----------
DEFAULT_SYSTEM = (
    "You are an expert YouTube script writer and finance educator for a faceless channel "
    "targeting Indian millennials (22-35). Tone: conversational, engaging, slightly dramatic, like a storyteller "
    "explaining money. Always: (1) open with a shocking hook in the first 30 seconds, (2) use simple language with "
    "no unexplained jargon, (3) add stock-footage scene directions in [brackets], (4) include a natural "
    "'This is not financial advice' line, (5) end with a subscribe CTA. Default target length: 1300-1500 words for "
    "a ~10-minute video. When asked for ideas, titles, hooks or thumbnails, be specific, punchy, and India-aware "
    "(₹, SIP, mutual funds, Groww, Zerodha references)."
)

@api.post("/chat/stream")
async def chat_stream(body: ChatRequest, user: dict = Depends(get_current_user)):
    settings = await db.user_settings.find_one({"user_id": user["id"]})
    if not settings or "anthropic_key_encrypted" not in settings:
        raise HTTPException(status_code=400, detail="No Anthropic API key configured. Add it in Settings.")
    try:
        key = decrypt_key(settings["anthropic_key_encrypted"])
    except Exception:
        raise HTTPException(status_code=500, detail="Failed to decrypt API key")

    client = Anthropic(api_key=key)
    system_prompt = body.system or DEFAULT_SYSTEM
    api_messages = [{"role": m.role, "content": m.content} for m in body.messages]

    # persist user message
    if body.video_id and body.messages:
        last_user = body.messages[-1]
        if last_user.role == "user":
            await db.chat_messages.insert_one({
                "id": str(uuid.uuid4()),
                "video_id": body.video_id,
                "role": "user",
                "content": last_user.content,
                "created_at": now_iso(),
            })

    def event_gen():
        full = ""
        try:
            with client.messages.stream(
                model="claude-sonnet-4-6",
                max_tokens=4096,
                system=system_prompt,
                messages=api_messages,
            ) as stream:
                for text in stream.text_stream:
                    full += text
                    yield text
        except Exception as e:
            err = f"\n\n[ERROR] {str(e)}"
            yield err
            full += err
        # persist assistant message
        if body.video_id:
            import asyncio
            try:
                loop = asyncio.new_event_loop()
                loop.run_until_complete(db.chat_messages.insert_one({
                    "id": str(uuid.uuid4()),
                    "video_id": body.video_id,
                    "role": "assistant",
                    "content": full,
                    "created_at": now_iso(),
                }))
                loop.close()
            except Exception as e:
                logger.error(f"failed to persist assistant msg: {e}")

    return StreamingResponse(event_gen(), media_type="text/plain")

# ---------- Uploads (Emergent object storage) ----------
ALLOWED_KIND = {
    "thumbnail": {"image/png", "image/jpeg", "image/webp"},
    "voiceover": {"audio/mpeg", "audio/mp3", "audio/wav", "audio/x-wav", "audio/mp4"},
}
EXT_MAP = {
    "image/png": "png", "image/jpeg": "jpg", "image/webp": "webp",
    "audio/mpeg": "mp3", "audio/mp3": "mp3", "audio/wav": "wav", "audio/x-wav": "wav", "audio/mp4": "m4a",
}

@api.post("/videos/{vid}/upload")
async def upload_for_video(vid: str, kind: str = Query(...), variant: str = Query("human"), file: UploadFile = File(...), user: dict = Depends(get_current_user)):
    if kind not in ALLOWED_KIND:
        raise HTTPException(400, "kind must be thumbnail or voiceover")
    if file.content_type not in ALLOWED_KIND[kind]:
        raise HTTPException(400, f"Unsupported content-type {file.content_type}")
    if variant not in ("human", "ai"):
        variant = "human"
    video = await db.videos.find_one({"id": vid})
    if not video:
        raise HTTPException(404, "Video not found")
    data = await file.read()
    if len(data) > 25 * 1024 * 1024:
        raise HTTPException(400, "File too large (max 25 MB)")
    ext = EXT_MAP.get(file.content_type, "bin")
    path = f"{APP_NAME}/uploads/{user['id']}/{vid}/{kind}-{uuid.uuid4()}.{ext}"
    try:
        result = put_object(path, data, file.content_type)
    except Exception as e:
        raise HTTPException(500, f"Upload failed: {e}")
    file_id = str(uuid.uuid4())
    await db.files.insert_one({
        "id": file_id,
        "video_id": vid,
        "kind": kind,
        "variant": variant if kind == "thumbnail" else None,
        "ctr": 0.0,
        "impressions": 0,
        "clicks": 0,
        "storage_path": result["path"],
        "original_filename": file.filename,
        "content_type": file.content_type,
        "size": result.get("size", len(data)),
        "is_deleted": False,
        "created_at": now_iso(),
    })
    field = "thumbnail_url" if kind == "thumbnail" else "voiceover_url"
    public_path = f"/api/files/{result['path']}"
    await db.videos.update_one({"id": vid}, {"$set": {field: public_path, "updated_at": now_iso()}})
    return {"id": file_id, "path": result["path"], "url": public_path, "variant": variant}

@api.get("/files/{path:path}")
async def serve_file(path: str, auth: str = Query(None), authorization: str = Header(None)):
    token = None
    if authorization and authorization.startswith("Bearer "):
        token = authorization[7:]
    elif auth:
        token = auth
    if not token:
        raise HTTPException(401, "Missing token")
    try:
        jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
    except Exception:
        raise HTTPException(401, "Invalid token")
    rec = await db.files.find_one({"storage_path": path, "is_deleted": False})
    if not rec:
        raise HTTPException(404, "File not found")
    data, ct = get_object(path)
    return Response(content=data, media_type=rec.get("content_type", ct))

# ---------- YouTube OAuth + Analytics API ----------
YT_OAUTH_SCOPES = [
    "https://www.googleapis.com/auth/yt-analytics.readonly",
    "https://www.googleapis.com/auth/youtube.readonly",
]

def _backend_origin() -> str:
    # CORS allows any; we need the public URL for redirect. Use FRONTEND_URL env or compute from request.
    return os.environ.get("PUBLIC_BACKEND_URL", "")

@api.get("/youtube/oauth/start")
async def yt_oauth_start(request: Request, user: dict = Depends(get_current_user)):
    client_id = await get_user_key(user["id"], "yt_oauth_client_id")
    client_secret = await get_user_key(user["id"], "yt_oauth_client_secret")
    if not client_id or not client_secret:
        raise HTTPException(400, "Add YouTube OAuth Client ID and Secret in Settings first")
    # Build redirect URI from request
    origin = str(request.base_url).rstrip("/")
    redirect_uri = f"{origin}/api/youtube/oauth/callback"
    flow = Flow.from_client_config(
        {"web": {
            "client_id": client_id,
            "client_secret": client_secret,
            "auth_uri": "https://accounts.google.com/o/oauth2/auth",
            "token_uri": "https://oauth2.googleapis.com/token",
            "redirect_uris": [redirect_uri],
        }},
        scopes=YT_OAUTH_SCOPES,
    )
    flow.redirect_uri = redirect_uri
    auth_url, state = flow.authorization_url(access_type="offline", prompt="consent", include_granted_scopes="true")
    await db.oauth_states.insert_one({"state": state, "user_id": user["id"], "redirect_uri": redirect_uri, "created_at": now_iso()})
    return {"auth_url": auth_url}

@api.get("/youtube/oauth/callback")
async def yt_oauth_callback(request: Request):
    params = dict(request.query_params)
    state = params.get("state")
    code = params.get("code")
    if not state or not code:
        return Response("Missing state/code", status_code=400)
    rec = await db.oauth_states.find_one({"state": state})
    if not rec:
        return Response("Invalid state", status_code=400)
    user_id = rec["user_id"]
    client_id = await get_user_key(user_id, "yt_oauth_client_id")
    client_secret = await get_user_key(user_id, "yt_oauth_client_secret")
    flow = Flow.from_client_config(
        {"web": {
            "client_id": client_id, "client_secret": client_secret,
            "auth_uri": "https://accounts.google.com/o/oauth2/auth",
            "token_uri": "https://oauth2.googleapis.com/token",
            "redirect_uris": [rec["redirect_uri"]],
        }},
        scopes=YT_OAUTH_SCOPES, state=state,
    )
    flow.redirect_uri = rec["redirect_uri"]
    try:
        flow.fetch_token(code=code)
    except Exception as e:
        return Response(f"OAuth failed: {e}", status_code=400)
    creds = flow.credentials
    refresh = creds.refresh_token
    if refresh:
        await db.user_settings.update_one(
            {"user_id": user_id},
            {"$set": {"yt_oauth_refresh_token_encrypted": encrypt_key(refresh), "yt_oauth_connected_at": now_iso()}},
            upsert=True,
        )
    await db.oauth_states.delete_one({"state": state})
    html = "<html><body style='font-family:sans-serif;padding:40px;text-align:center'><h2 style='color:#00594C'>YouTube connected ✓</h2><p>You can close this tab and return to Wealth Studio.</p><script>setTimeout(()=>window.close(),1500)</script></body></html>"
    return Response(content=html, media_type="text/html")

@api.get("/youtube/oauth/status")
async def yt_oauth_status(user: dict = Depends(get_current_user)):
    doc = await db.user_settings.find_one({"user_id": user["id"]})
    return {"connected": bool((doc or {}).get("yt_oauth_refresh_token_encrypted")), "connected_at": (doc or {}).get("yt_oauth_connected_at")}

@api.delete("/youtube/oauth")
async def yt_oauth_disconnect(user: dict = Depends(get_current_user)):
    await db.user_settings.update_one({"user_id": user["id"]}, {"$unset": {"yt_oauth_refresh_token_encrypted": "", "yt_oauth_connected_at": ""}})
    return {"status": "disconnected"}

async def _get_yt_credentials(user_id: str):
    doc = await db.user_settings.find_one({"user_id": user_id})
    enc = (doc or {}).get("yt_oauth_refresh_token_encrypted")
    if not enc:
        return None
    refresh_token = decrypt_key(enc)
    client_id = await get_user_key(user_id, "yt_oauth_client_id")
    client_secret = await get_user_key(user_id, "yt_oauth_client_secret")
    if not client_id or not client_secret:
        return None
    creds = Credentials(
        token=None, refresh_token=refresh_token,
        token_uri="https://oauth2.googleapis.com/token",
        client_id=client_id, client_secret=client_secret,
        scopes=YT_OAUTH_SCOPES,
    )
    await asyncio.to_thread(creds.refresh, GoogleAuthRequest())
    return creds

@api.get("/videos/{vid}/yt-analytics")
async def yt_analytics(vid: str, user: dict = Depends(get_current_user)):
    video = await db.videos.find_one({"id": vid}, {"_id": 0})
    if not video:
        raise HTTPException(404, "Video not found")
    yt_id = _extract_video_id(video.get("youtube_url") or "")
    if not yt_id:
        raise HTTPException(400, "Set the YouTube URL on this card first")
    creds = await _get_yt_credentials(user["id"])
    if not creds:
        raise HTTPException(400, "Connect YouTube in Settings → YouTube Analytics first")

    def fetch():
        ya = gbuild("youtubeAnalytics", "v2", credentials=creds, cache_discovery=False)
        end = datetime.now(timezone.utc).date().isoformat()
        start = "2020-01-01"
        retention = ya.reports().query(
            ids="channel==MINE", startDate=start, endDate=end,
            metrics="audienceWatchRatio,relativeRetentionPerformance",
            dimensions="elapsedVideoTimeRatio", filters=f"video=={yt_id}",
        ).execute()
        sources = ya.reports().query(
            ids="channel==MINE", startDate=start, endDate=end,
            metrics="views,estimatedMinutesWatched,averageViewDuration",
            dimensions="insightTrafficSourceType", filters=f"video=={yt_id}",
            sort="-views",
        ).execute()
        summary = ya.reports().query(
            ids="channel==MINE", startDate=start, endDate=end,
            metrics="views,estimatedMinutesWatched,averageViewDuration,subscribersGained,likes,comments,shares",
            filters=f"video=={yt_id}",
        ).execute()
        return retention, sources, summary

    try:
        retention, sources, summary = await asyncio.to_thread(fetch)
    except Exception as e:
        raise HTTPException(500, f"YouTube Analytics error: {e}")

    def rows(report):
        headers = [h["name"] for h in report.get("columnHeaders", [])]
        return [dict(zip(headers, r)) for r in report.get("rows", [])]

    return {
        "retention_curve": rows(retention),  # [{elapsedVideoTimeRatio, audienceWatchRatio, relativeRetentionPerformance}]
        "traffic_sources": rows(sources),    # [{insightTrafficSourceType, views, ...}]
        "summary": (rows(summary)[0] if summary.get("rows") else {}),
    }

# ---------- Thumbnail A/B tracker ----------
class ThumbStatsUpdate(BaseModel):
    impressions: Optional[int] = None
    clicks: Optional[int] = None
    ctr: Optional[float] = None

@api.get("/videos/{vid}/thumbnails")
async def list_thumbnails(vid: str, user: dict = Depends(get_current_user)):
    docs = await db.files.find(
        {"video_id": vid, "kind": "thumbnail", "is_deleted": False},
        {"_id": 0},
    ).sort("created_at", 1).to_list(50)
    for d in docs:
        d["url"] = f"/api/files/{d['storage_path']}"
    return docs

@api.patch("/thumbnails/{file_id}")
async def update_thumb_stats(file_id: str, body: ThumbStatsUpdate, user: dict = Depends(get_current_user)):
    updates = {k: v for k, v in body.model_dump().items() if v is not None}
    if "impressions" in updates and "clicks" in updates and updates["impressions"] > 0:
        updates["ctr"] = round(updates["clicks"] / updates["impressions"] * 100, 2)
    if updates:
        await db.files.update_one({"id": file_id, "kind": "thumbnail"}, {"$set": updates})
    doc = await db.files.find_one({"id": file_id}, {"_id": 0})
    if not doc:
        raise HTTPException(404, "Thumbnail not found")
    doc["url"] = f"/api/files/{doc['storage_path']}"
    return doc

@api.delete("/thumbnails/{file_id}")
async def delete_thumb(file_id: str, user: dict = Depends(get_current_user)):
    await db.files.update_one({"id": file_id}, {"$set": {"is_deleted": True}})
    return {"status": "deleted"}

# ---------- Script export ----------
@api.get("/videos/{vid}/export")
async def export_script(vid: str, fmt: str = Query("txt"), user: dict = Depends(get_current_user)):
    v = await db.videos.find_one({"id": vid}, {"_id": 0})
    if not v:
        raise HTTPException(404, "Video not found")
    script = v.get("script") or ""
    if not script.strip():
        raise HTTPException(400, "Script is empty — write one in the Claude chat first.")
    safe_title = "".join(c if c.isalnum() or c in " -_" else "_" for c in v["title"])[:60].strip() or "script"

    if fmt == "txt":
        content = f"{v['title']}\n{'='*len(v['title'])}\n\n{script}\n"
        return StreamingResponse(
            io.BytesIO(content.encode("utf-8")),
            media_type="text/plain",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.txt"'},
        )
    if fmt == "docx":
        from docx import Document
        doc = Document()
        doc.add_heading(v["title"], level=1)
        if v.get("hook"):
            p = doc.add_paragraph()
            p.add_run("Hook: ").bold = True
            p.add_run(v["hook"])
        doc.add_paragraph("")
        for para in script.split("\n"):
            doc.add_paragraph(para)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.docx"'},
        )
    raise HTTPException(400, "fmt must be txt or docx")

# ---------- Seed ----------
SEED_IDEAS = [
    # Personal Finance India
    ("Why 90% of Indians Die Broke (And How to Not Be One of Them)", "personal_finance"),
    ("The Real Reason Your Salary Is Never Enough", "personal_finance"),
    ("The 50-30-20 Rule Indians Are Doing Wrong", "personal_finance"),
    ("Biggest Money Mistakes Indians Make in Their 20s", "personal_finance"),
    ("Why Gold Is NOT What Rich Indians Buy", "personal_finance"),
    ("The Truth About LIC Policies — Are You Getting Scammed?", "personal_finance"),
    ("How to Build an Emergency Fund That Actually Works", "personal_finance"),
    ("5 Things Rich Indians Do Differently with Their Salary", "personal_finance"),
    ("Hidden Costs Nobody Warns You About in India", "personal_finance"),
    ("How Much Should You Really Be Saving Every Month?", "personal_finance"),
    # Investing
    ("Mutual Funds Explained in 10 Minutes (For Absolute Beginners)", "investing"),
    ("SIP vs Lump Sum: Which One Makes You More Money?", "investing"),
    ("What is an Index Fund? Why Warren Buffett Swears By It", "investing"),
    ("How to Start Investing with Just ₹500/month", "investing"),
    ("Nifty 50 vs Sensex — Finally Explained Simply", "investing"),
    ("ELSS vs PPF: Which Tax-Saving Investment is Better?", "investing"),
    ("Compounding: The Most Powerful Concept in Finance", "investing"),
    ("Groww vs Zerodha vs Upstox — Which One Should You Pick?", "investing"),
    ("What is a Demat Account? Step-by-Step How to Open", "investing"),
    ("Active vs Passive Mutual Funds: Which Wins Long Term?", "investing"),
    # Psychology
    ("Why Smart People Make Dumb Money Decisions", "psychology"),
    ("The Psychology of Why You Can't Stop Spending", "psychology"),
    ("How the Rich Think About Money Differently", "psychology"),
    ("7 Financial Habits That Changed My Life", "psychology"),
    ("The Lifestyle Inflation Trap — Why Raises Make You Poorer", "psychology"),
    ("Why Most People Never Become Wealthy (The Real Reason)", "psychology"),
    ("The One Money Mistake That Keeps People Broke Forever", "psychology"),
    ("What School Never Taught You About Money", "psychology"),
    ("How to Think About Money Like a Millionaire", "psychology"),
    ("The 3 Types of Income — Only One Makes You Rich", "psychology"),
    # Case Studies
    ("How Zepto Went from Zero to ₹10,000 Crore in 2 Years", "case_study"),
    ("Why Byju's Failed: India's Biggest Startup Collapse", "case_study"),
    ("How Zomato Built a ₹1 Lakh Crore Company From Scratch", "case_study"),
    ("The Rise and Fall of Kingfisher Airlines", "case_study"),
    ("How Jio Changed India's Economy Forever", "case_study"),
    ("Paytm's Billion Dollar Mistake — What Really Happened", "case_study"),
    ("How Amazon Entered India and Disrupted Everything", "case_study"),
    ("CRED — How a Loss-Making App Raised Billions", "case_study"),
    ("Why Flipkart Was Sold to Walmart — The Inside Story", "case_study"),
    ("How Tata Built India's Most Trusted Empire", "case_study"),
    # Side Hustles
    ("I Made ₹50,000 in One Month With No Investment", "side_hustle"),
    ("5 Side Hustles That Actually Pay in India in 2026", "side_hustle"),
    ("Passive Income Methods That Actually Work (Tested)", "side_hustle"),
    ("How to Quit Your Job in 12 Months (Real Plan)", "side_hustle"),
    ("Credit Card Hacks Indians Don't Know About", "side_hustle"),
    ("Best Credit Cards in India in 2026 (Ranked)", "side_hustle"),
    ("Home Loan vs Rent: What's Actually Better in India?", "side_hustle"),
    ("How to File ITR Yourself in 30 Minutes", "side_hustle"),
    ("What is CIBIL Score? How to Improve It Fast", "side_hustle"),
    ("How to Negotiate a Higher Salary in India (Scripts Included)", "side_hustle"),
]

async def seed_admin_and_ideas():
    # Admin user
    user = await db.users.find_one({"email": ADMIN_EMAIL})
    if not user:
        await db.users.insert_one({
            "id": str(uuid.uuid4()),
            "email": ADMIN_EMAIL,
            "password_hash": hash_password(ADMIN_PASSWORD),
            "name": "Creator",
            "created_at": now_iso(),
        })
        logger.info("Admin user seeded.")
    else:
        # rotate password if .env changed
        if not verify_password(ADMIN_PASSWORD, user["password_hash"]):
            await db.users.update_one({"email": ADMIN_EMAIL}, {"$set": {"password_hash": hash_password(ADMIN_PASSWORD)}})
            logger.info("Admin password updated from .env.")

    # Seed ideas only if empty
    count = await db.ideas.count_documents({})
    if count == 0:
        docs = []
        for title, niche in SEED_IDEAS:
            docs.append(Idea(title=title, sub_niche=niche, rating=7).model_dump())
        await db.ideas.insert_many(docs)
        logger.info(f"Seeded {len(docs)} ideas.")

# ---------- Wire up ----------
app.include_router(api)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("startup")
async def on_start():
    await seed_admin_and_ideas()
    try:
        init_storage()
        logger.info("Object storage initialized.")
    except Exception as e:
        logger.warning(f"Storage init deferred: {e}")

@app.on_event("shutdown")
async def on_stop():
    mongo_client.close()
